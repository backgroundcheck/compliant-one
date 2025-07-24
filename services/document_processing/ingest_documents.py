import os
import pandas as pd
import sqlite3
import uuid
from typing import Dict, Any, List
from langchain.docstore.document import Document
from langchain.vectorstores import Chroma
from langchain_cohere import CohereEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
import fitz  # PyMuPDF
from pdf2image import convert_from_path
import pytesseract
import docx
import json
import xml.etree.ElementTree as ET
from tqdm import tqdm

# Configuration
INPUT_DIR = "./downloaded_pdfs"  # Directory containing PDF files to ingest
COLLECTION_NAME = "document_collection"
COHERE_API_KEY = "ONVwImVb9P2ds0G16XGocuFSuL7o8C0eij6xZRd8"
EMBEDDING_MODEL = "embed-english-v3.0"
DB_PATH = "documents.db"

# Initialize SQLite database
def init_sqlite_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            document_id TEXT PRIMARY KEY,
            filename TEXT,
            filetype TEXT,
            site TEXT,
            publisher TEXT,
            dataset TEXT,
            officialtitle TEXT,
            url TEXT,
            exists_flag BOOLEAN,
            digital BOOLEAN,
            public BOOLEAN,
            online BOOLEAN,
            free BOOLEAN,
            machinereadable BOOLEAN,
            bulk BOOLEAN,
            openlicence BOOLEAN,
            uptodate BOOLEAN,
            extracted_title TEXT,
            extracted_author TEXT
        )
    """)
    conn.commit()
    conn.close()

# Convert CSV values to boolean or None
def to_boolean(value: Any) -> bool:
    if pd.isna(value):
        return None
    value = str(value).strip().lower()
    if value == "yes":
        return True
    if value == "no":
        return False
    return None

def row_to_markdown(row: pd.Series) -> str:
    """Convert a pandas row to a Markdown list."""
    lines = [f"- **{col}**: {row[col]}" for col in row.index if pd.notna(row[col])]
    return "\n".join(lines)

# Extract text and metadata from CSV
def process_csv(file_path: str) -> List[Dict[str, Any]]:
    df = pd.read_csv(file_path)
    documents = []
    for _, row in df.iterrows():
        text = row_to_markdown(row)
        metadata = {
            "document_id": str(uuid.uuid4()),
            "filename": os.path.basename(file_path),
            "filetype": "csv",
            "site": str(row["site"]) if pd.notna(row["site"]) else None,
            "publisher": str(row["publisher"]) if pd.notna(row["publisher"]) else None,
            "dataset": str(row["dataset"]) if pd.notna(row["dataset"]) else None,
            "officialtitle": str(row["officialtitle"]) if pd.notna(row["officialtitle"]) else None,
            "url": str(row["url"]) if pd.notna(row["url"]) else None,
            "exists_flag": to_boolean(row["exists"]),
            "digital": to_boolean(row["digital"]),
            "public": to_boolean(row["public"]),
            "online": to_boolean(row["online"]),
            "free": to_boolean(row["free"]),
            "machinereadable": to_boolean(row["machinereadable"]),
            "bulk": to_boolean(row["bulk"]),
            "openlicence": to_boolean(row["openlicence"]),
            "uptodate": to_boolean(row["uptodate"]),
            "extracted_title": None,
            "extracted_author": None
        }
        # Create text representation for embedding
        text = (
            f"The dataset titled {metadata['officialtitle'] or 'Untitled'} "
            f"published by {metadata['publisher'] or 'Unknown'} on {metadata['site'] or 'Unknown'} "
            f"is {'digital' if metadata['digital'] else 'not digital'}, "
            f"{'public' if metadata['public'] else 'not public'}, "
            f"{'free' if metadata['free'] else 'not free'}, "
            f"{'bulk-downloadable' if metadata['bulk'] else 'not bulk-downloadable'}, "
            f"licensed {metadata['url'] or 'with no specified license'}."
        )
        documents.append({"text": text, "metadata": metadata})
    return documents

# Extract text and metadata from PDF
def process_pdf(file_path: str) -> List[Dict[str, Any]]:
    documents = []
    try:
        doc = fitz.open(file_path)
        full_text = ""
        for page in doc:
            text = page.get_text()
            if not text.strip():
                # Fallback to OCR if no text found
                images = convert_from_path(file_path, first_page=page.number+1, last_page=page.number+1)
                for image in images:
                    text += pytesseract.image_to_string(image)
            full_text += text
        metadata = {
            "document_id": str(uuid.uuid4()),
            "filename": os.path.basename(file_path),
            "filetype": "pdf",
            "site": None,
            "publisher": None,
            "dataset": None,
            "officialtitle": None,
            "url": None,
            "exists_flag": None,
            "digital": None,
            "public": None,
            "online": None,
            "free": None,
            "machinereadable": None,
            "bulk": None,
            "openlicence": None,
            "uptodate": None,
            "extracted_title": doc.metadata.get("title"),
            "extracted_author": doc.metadata.get("author")
        }
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            length_function=len
        )
        chunks = text_splitter.split_text(full_text)
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata["document_id"] = str(uuid.uuid4())
            chunk_metadata["chunk_index"] = i
            documents.append({"text": chunk, "metadata": chunk_metadata})
    except Exception as e:
        print(f"[WARN] Skipping file (not a valid PDF): {file_path} | Reason: {e}")
    return documents

# Extract text and metadata from Word document
def process_docx(file_path: str) -> List[Dict[str, Any]]:
    doc = docx.Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text])
    metadata = {
        "document_id": str(uuid.uuid4()),
        "filename": os.path.basename(file_path),
        "filetype": "docx",
        "site": None,
        "publisher": None,
        "dataset": None,
        "officialtitle": None,
        "url": None,
        "exists_flag": None,
        "digital": None,
        "public": None,
        "online": None,
        "free": None,
        "machinereadable": None,
        "bulk": None,
        "openlicence": None,
        "uptodate": None,
        "extracted_title": None,
        "extracted_author": None
    }
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        length_function=len
    )
    chunks = text_splitter.split_text(full_text)
    documents = []
    for i, chunk in enumerate(chunks):
        chunk_metadata = metadata.copy()
        chunk_metadata["document_id"] = str(uuid.uuid4())
        chunk_metadata["chunk_index"] = i
        documents.append({"text": chunk, "metadata": chunk_metadata})
    return documents

def process_excel(file_path: str) -> List[Dict[str, Any]]:
    try:
        df = pd.read_excel(file_path, engine='openpyxl' if file_path.endswith('.xlsx') else 'xlrd')
    except Exception as e:
        print(f"[WARN] Skipping file (not a valid Excel): {file_path} | Reason: {e}")
        return []
    documents = []
    for _, row in df.iterrows():
        text = row_to_markdown(row)
        metadata = {
            "document_id": str(uuid.uuid4()),
            "filename": os.path.basename(file_path),
            "filetype": "xlsx" if file_path.endswith('.xlsx') else "xls",
            "extracted_title": None,
            "extracted_author": None
        }
        documents.append({"text": text, "metadata": metadata})
    return documents


def process_json(file_path: str) -> List[Dict[str, Any]]:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        text = "```json\n" + json.dumps(data, indent=2) + "\n```"
    except Exception as e:
        print(f"[WARN] Skipping file (not a valid JSON): {file_path} | Reason: {e}")
        return []
    metadata = {
        "document_id": str(uuid.uuid4()),
        "filename": os.path.basename(file_path),
        "filetype": "json",
        "extracted_title": None,
        "extracted_author": None
    }
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    documents = []
    for i, chunk in enumerate(chunks):
        chunk_metadata = metadata.copy()
        chunk_metadata["document_id"] = str(uuid.uuid4())
        chunk_metadata["chunk_index"] = i
        documents.append({"text": chunk, "metadata": chunk_metadata})
    return documents


def process_xml(file_path: str) -> List[Dict[str, Any]]:
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        text = "```xml\n" + ET.tostring(root, encoding='unicode') + "\n```"
    except Exception as e:
        print(f"[WARN] Skipping file (not a valid XML): {file_path} | Reason: {e}")
        return []
    metadata = {
        "document_id": str(uuid.uuid4()),
        "filename": os.path.basename(file_path),
        "filetype": "xml",
        "extracted_title": None,
        "extracted_author": None
    }
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    documents = []
    for i, chunk in enumerate(chunks):
        chunk_metadata = metadata.copy()
        chunk_metadata["document_id"] = str(uuid.uuid4())
        chunk_metadata["chunk_index"] = i
        documents.append({"text": chunk, "metadata": chunk_metadata})
    return documents

# Save metadata to SQLite
def save_to_sqlite(documents: List[Dict[str, Any]]):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    for doc in documents:
        cursor.execute("""
            INSERT INTO documents (
                document_id, filename, filetype, site, publisher, dataset, officialtitle, url,
                exists_flag, digital, public, online, free, machinereadable, bulk, openlicence, uptodate,
                extracted_title, extracted_author
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            doc["metadata"]["document_id"],
            doc["metadata"]["filename"],
            doc["metadata"]["filetype"],
            doc["metadata"]["site"],
            doc["metadata"]["publisher"],
            doc["metadata"]["dataset"],
            doc["metadata"]["officialtitle"],
            doc["metadata"]["url"],
            doc["metadata"]["exists_flag"],
            doc["metadata"]["digital"],
            doc["metadata"]["public"],
            doc["metadata"]["online"],
            doc["metadata"]["free"],
            doc["metadata"]["machinereadable"],
            doc["metadata"]["bulk"],
            doc["metadata"]["openlicence"],
            doc["metadata"]["uptodate"],
            doc["metadata"]["extracted_title"],
            doc["metadata"]["extracted_author"]
        ))
    conn.commit()
    conn.close()

# Main ingestion process
def ingest_documents():
    init_sqlite_db()
    documents = []
    files = os.listdir(INPUT_DIR)
    print(f"Processing {len(files)} files in {INPUT_DIR}...")
    for filename in tqdm(files, desc="Extracting text & metadata"):
        file_path = os.path.join(INPUT_DIR, filename)
        if filename.endswith(".csv"):
            docs = process_csv(file_path)
        elif filename.endswith(".pdf"):
            docs = process_pdf(file_path)
        elif filename.endswith(".docx"):
            docs = process_docx(file_path)
        elif filename.endswith(".xlsx") or filename.endswith(".xls"):
            docs = process_excel(file_path)
        elif filename.endswith(".json"):
            docs = process_json(file_path)
        elif filename.endswith(".xml"):
            docs = process_xml(file_path)
        else:
            continue  # Skip unsupported file types
        documents.extend(docs)
    # Save metadata to SQLite
    print("Saving metadata to SQLite...")
    save_to_sqlite(documents)
    # Create LangChain Documents for vector store
    print("Creating LangChain documents...")
    langchain_docs = [
        Document(page_content=doc["text"], metadata=doc["metadata"])
        for doc in tqdm(documents, desc="Preparing vector docs")
    ]
    # Initialize Cohere embeddings and Chroma vector store
    print("Embedding and adding to Chroma vector store...")
    embeddings = CohereEmbeddings(model=EMBEDDING_MODEL)
    vector_store = Chroma(
        collection_name=COLLECTION_NAME,
        embedding_function=embeddings,
        persist_directory="./chroma_db"
    )
    vector_store.add_documents(langchain_docs)
    vector_store.persist()
    print(f"Successfully ingested {len(documents)} documents into {COLLECTION_NAME} collection.")

if __name__ == "__main__":
    ingest_documents()