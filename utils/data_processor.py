"""
Data Processing Utilities for Compliant.one Admin Dashboard
Handles conversion of various file formats to structured data
"""

import pandas as pd
import json
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional, Union
import csv
import re
from datetime import datetime

# Import for different file types
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import PyPDF2 # type: ignore
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class DataProcessor:
    """Process different file formats and extract structured data"""
    
    def __init__(self):
        self.supported_formats = {
            'csv': self.process_csv,
            'xml': self.process_xml,
            'json': self.process_json,
            'txt': self.process_text,
            'html': self.process_html if BS4_AVAILABLE else None,
            'xlsx': self.process_excel if EXCEL_AVAILABLE else None,
            'xls': self.process_excel if EXCEL_AVAILABLE else None,
            'docx': self.process_docx if DOCX_AVAILABLE else None,
            'pdf': self.process_pdf if PDF_AVAILABLE else None
        }
    
    def process_file(self, file_path: Path, file_type: str) -> Dict:
        """Process file based on its type"""
        if file_type not in self.supported_formats:
            return {
                'status': 'error',
                'message': f'Unsupported file type: {file_type}',
                'entities': []
            }
        
        processor = self.supported_formats[file_type]
        if processor is None:
            return {
                'status': 'error',
                'message': f'Required library not available for {file_type} files',
                'entities': []
            }
        
        try:
            return processor(file_path)
        except Exception as e:
            return {
                'status': 'error',
                'message': f'Error processing {file_type} file: {str(e)}',
                'entities': []
            }
    
    def process_csv(self, file_path: Path) -> Dict:
        """Process CSV file"""
        entities = []
        
        try:
            df = pd.read_csv(file_path)
            
            for index, row in df.iterrows():
                for column in df.columns:
                    value = str(row[column]) if pd.notna(row[column]) else ''
                    if value.strip():
                        entities.append({
                            'type': 'tabular_data',
                            'name': value,
                            'source_row': index + 1,
                            'source_column': column,
                            'metadata': {
                                'file_type': 'csv',
                                'row': index + 1,
                                'column': column,
                                'total_rows': len(df),
                                'total_columns': len(df.columns)
                            }
                        })
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from CSV with {len(df)} rows',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'CSV processing error: {str(e)}',
                'entities': []
            }
    
    def process_xml(self, file_path: Path) -> Dict:
        """Process XML file"""
        entities = []
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            def extract_entities(element, path=""):
                current_path = f"{path}/{element.tag}" if path else element.tag
                
                # Extract text content
                if element.text and element.text.strip():
                    entities.append({
                        'type': 'xml_element',
                        'name': element.text.strip(),
                        'element_tag': element.tag,
                        'element_path': current_path,
                        'metadata': {
                            'file_type': 'xml',
                            'tag': element.tag,
                            'path': current_path,
                            'attributes': element.attrib
                        }
                    })
                
                # Extract attributes
                for attr_name, attr_value in element.attrib.items():
                    if attr_value.strip():
                        entities.append({
                            'type': 'xml_attribute',
                            'name': attr_value,
                            'element_tag': element.tag,
                            'attribute_name': attr_name,
                            'metadata': {
                                'file_type': 'xml',
                                'tag': element.tag,
                                'attribute': attr_name,
                                'path': current_path
                            }
                        })
                
                # Process children
                for child in element:
                    extract_entities(child, current_path)
            
            extract_entities(root)
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from XML',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'XML processing error: {str(e)}',
                'entities': []
            }
    
    def process_json(self, file_path: Path) -> Dict:
        """Process JSON file"""
        entities = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            def extract_from_json(obj, path=""):
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        current_path = f"{path}.{key}" if path else key
                        
                        if isinstance(value, (str, int, float)) and str(value).strip():
                            entities.append({
                                'type': 'json_value',
                                'name': str(value),
                                'key': key,
                                'path': current_path,
                                'metadata': {
                                    'file_type': 'json',
                                    'key': key,
                                    'path': current_path,
                                    'value_type': type(value).__name__
                                }
                            })
                        elif isinstance(value, (dict, list)):
                            extract_from_json(value, current_path)
                
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        current_path = f"{path}[{i}]" if path else f"[{i}]"
                        extract_from_json(item, current_path)
            
            extract_from_json(data)
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from JSON',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'JSON processing error: {str(e)}',
                'entities': []
            }
    
    def process_text(self, file_path: Path) -> Dict:
        """Process plain text file"""
        entities = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            for line_num, line in enumerate(lines, 1):
                line = line.strip()
                if line:
                    # Extract potential entities using basic patterns
                    # Email addresses
                    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', line)
                    for email in emails:
                        entities.append({
                            'type': 'email',
                            'name': email,
                            'line_number': line_num,
                            'metadata': {
                                'file_type': 'text',
                                'entity_type': 'email',
                                'line': line_num,
                                'context': line[:100]
                            }
                        })
                    
                    # Phone numbers (basic pattern)
                    phones = re.findall(r'\b(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', line)
                    for phone in phones:
                        entities.append({
                            'type': 'phone',
                            'name': phone,
                            'line_number': line_num,
                            'metadata': {
                                'file_type': 'text',
                                'entity_type': 'phone',
                                'line': line_num,
                                'context': line[:100]
                            }
                        })
                    
                    # Generic text entity
                    entities.append({
                        'type': 'text_line',
                        'name': line,
                        'line_number': line_num,
                        'metadata': {
                            'file_type': 'text',
                            'entity_type': 'text_line',
                            'line': line_num,
                            'length': len(line)
                        }
                    })
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from text file',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'Text processing error: {str(e)}',
                'entities': []
            }
    
    def process_html(self, file_path: Path) -> Dict:
        """Process HTML file"""
        if not BS4_AVAILABLE:
            return {
                'status': 'error',
                'message': 'BeautifulSoup4 not available for HTML processing',
                'entities': []
            }
        
        entities = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
            
            # Extract text from various HTML elements
            text_tags = ['p', 'div', 'span', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'td', 'th']
            
            for tag_name in text_tags:
                elements = soup.find_all(tag_name)
                for i, element in enumerate(elements):
                    text = element.get_text(strip=True)
                    if text:
                        entities.append({
                            'type': 'html_text',
                            'name': text,
                            'element_tag': tag_name,
                            'element_index': i,
                            'metadata': {
                                'file_type': 'html',
                                'tag': tag_name,
                                'index': i,
                                'attributes': element.attrs
                            }
                        })
            
            # Extract links
            links = soup.find_all('a', href=True)
            for i, link in enumerate(links):
                href = link['href']
                text = link.get_text(strip=True)
                entities.append({
                    'type': 'html_link',
                    'name': f"{text} ({href})" if text else href,
                    'url': href,
                    'link_text': text,
                    'metadata': {
                        'file_type': 'html',
                        'entity_type': 'link',
                        'href': href,
                        'text': text
                    }
                })
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from HTML',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'HTML processing error: {str(e)}',
                'entities': []
            }
    
    def process_excel(self, file_path: Path) -> Dict:
        """Process Excel file"""
        if not EXCEL_AVAILABLE:
            return {
                'status': 'error',
                'message': 'openpyxl not available for Excel processing',
                'entities': []
            }
        
        entities = []
        
        try:
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                for index, row in df.iterrows():
                    for column in df.columns:
                        value = str(row[column]) if pd.notna(row[column]) else ''
                        if value.strip():
                            entities.append({
                                'type': 'excel_cell',
                                'name': value,
                                'sheet': sheet_name,
                                'row': index + 1,
                                'column': column,
                                'metadata': {
                                    'file_type': 'excel',
                                    'sheet': sheet_name,
                                    'row': index + 1,
                                    'column': column,
                                    'total_sheets': len(excel_file.sheet_names)
                                }
                            })
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from {len(excel_file.sheet_names)} Excel sheets',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'Excel processing error: {str(e)}',
                'entities': []
            }
    
    def process_docx(self, file_path: Path) -> Dict:
        """Process Word document"""
        if not DOCX_AVAILABLE:
            return {
                'status': 'error',
                'message': 'python-docx not available for DOCX processing',
                'entities': []
            }
        
        entities = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract from paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    entities.append({
                        'type': 'docx_paragraph',
                        'name': text,
                        'paragraph_index': i,
                        'metadata': {
                            'file_type': 'docx',
                            'element_type': 'paragraph',
                            'index': i
                        }
                    })
            
            # Extract from tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if text:
                            entities.append({
                                'type': 'docx_table_cell',
                                'name': text,
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'cell_index': cell_idx,
                                'metadata': {
                                    'file_type': 'docx',
                                    'element_type': 'table_cell',
                                    'table': table_idx,
                                    'row': row_idx,
                                    'cell': cell_idx
                                }
                            })
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from DOCX',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'DOCX processing error: {str(e)}',
                'entities': []
            }
    
    def process_pdf(self, file_path: Path) -> Dict:
        """Process PDF file"""
        if not PDF_AVAILABLE:
            return {
                'status': 'error',
                'message': 'PyPDF2 not available for PDF processing',
                'entities': []
            }
        
        entities = []
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            # Split into lines and process
                            lines = text.split('\n')
                            for line_num, line in enumerate(lines):
                                line = line.strip()
                                if line:
                                    entities.append({
                                        'type': 'pdf_text',
                                        'name': line,
                                        'page': page_num + 1,
                                        'line': line_num + 1,
                                        'metadata': {
                                            'file_type': 'pdf',
                                            'page': page_num + 1,
                                            'line': line_num + 1,
                                            'total_pages': len(pdf_reader.pages)
                                        }
                                    })
                    except Exception as e:
                        # Skip pages that can't be processed
                        continue
            
            return {
                'status': 'success',
                'message': f'Processed {len(entities)} entities from PDF with {len(pdf_reader.pages)} pages',
                'entities': entities
            }
            
        except Exception as e:
            return {
                'status': 'error',
                'message': f'PDF processing error: {str(e)}',
                'entities': []
            }

# Entity categorization and risk scoring
class EntityClassifier:
    """Classify extracted entities and assign risk scores"""
    
    def __init__(self):
        self.high_risk_patterns = [
            r'terror', r'money\s*laundering', r'fraud', r'corruption',
            r'sanctions?', r'embargo', r'prohibited', r'blacklist'
        ]
        
        self.medium_risk_patterns = [
            r'investigation', r'suspicious', r'violation', r'penalty',
            r'enforcement', r'compliance', r'regulatory'
        ]
        
        self.entity_types = {
            'person': r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b',
            'organization': r'\b[A-Z][A-Z\s&]+(?:INC|LLC|LTD|CORP|COMPANY)\b',
            'location': r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2,3}\b',
            'identifier': r'\b[A-Z0-9]{6,}\b'
        }
    
    def classify_entity(self, entity: Dict) -> Dict:
        """Classify entity and assign risk score"""
        name = entity.get('name', '').lower()
        
        # Calculate risk score
        risk_score = self._calculate_risk_score(name)
        
        # Determine entity type
        entity_type = self._determine_entity_type(entity.get('name', ''))
        
        # Add classification
        entity['risk_score'] = risk_score
        entity['entity_type'] = entity_type
        entity['risk_level'] = self._get_risk_level(risk_score)
        
        return entity
    
    def _calculate_risk_score(self, text: str) -> float:
        """Calculate risk score based on text content"""
        score = 0.0
        
        # Check for high-risk patterns
        for pattern in self.high_risk_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                score += 0.8
        
        # Check for medium-risk patterns
        for pattern in self.medium_risk_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                score += 0.4
        
        # Normalize score to 0-1 range
        return min(score, 1.0)
    
    def _determine_entity_type(self, text: str) -> str:
        """Determine the type of entity"""
        for entity_type, pattern in self.entity_types.items():
            if re.search(pattern, text):
                return entity_type
        
        return 'unknown'
    
    def _get_risk_level(self, risk_score: float) -> str:
        """Convert risk score to risk level"""
        if risk_score >= 0.7:
            return 'high'
        elif risk_score >= 0.4:
            return 'medium'
        else:
            return 'low'
